import pandas as pd
from docx import Document
import os

def criar_ficha_projeto(linha):
    """
    Função responsável por gerar o documento Word (Ficha de Projeto) 
    para uma linha específica (projeto) do dataframe.
    """
    doc = Document()
    
    # Cabeçalho Principal
    doc.add_heading('EMBRAPII', level=0)
    doc.add_heading('Ficha de Projeto', level=1)
    
    # Função auxiliar para adicionar rótulo em negrito e valor normal na mesma linha
    def adicionar_campo(rotulo, valor):
        p = doc.add_paragraph()
        p.add_run(f"{rotulo}: ").bold = True
        # Converte o valor para string para evitar erros com números ou datas
        p.add_run(str(valor))

    # 1. Identificação do Projeto
    doc.add_heading('1. Identificação do Projeto', level=2)
    adicionar_campo('Código do Projeto', linha['Projeto_cód'])
    adicionar_campo('Empresa Contratante (cód.)', linha['Empresas Contratantes_cód'])
    adicionar_campo('Status do Projeto', linha['Projeto_status'])
    adicionar_campo('Tipo de Projeto', linha['Tipo de Projeto'])
    
    # 2. Datas e Duração
    doc.add_heading('2. Datas e Duração', level=2)
    adicionar_campo('Data do Contrato', linha['Data_Contrato'])
    adicionar_campo('Data de Início', linha['Data_Início'])
    adicionar_campo('Data de Término', linha['Data_Término'])
    adicionar_campo('Duração (dias)', linha['Data_Duração_dias'])
    
    # 3. Unidade Embrapii
    doc.add_heading('3. Unidade Embrapii', level=2)
    adicionar_campo('Código da Unidade', linha['Unidade Embrapii_cód'])
    adicionar_campo('Status da Unidade', linha['Unidade Embrapii_status'])
    adicionar_campo('Região', linha['Unidade Embrapii_Região'])
    adicionar_campo('UF', linha['Unidade Embrapii_UF'])
    adicionar_campo('Tipo de Instituição', linha['Unidade Embrapii_Tipo de Instituição'])
    
    # 4. Informações Financeiras
    doc.add_heading('4. Informações Financeiras', level=2)
    adicionar_campo('Fonte de Recursos', linha['Fin_Fonte de Recursos'])
    adicionar_campo('Valor Aporte Embrapii (R$)', linha['Fin_Valor aporte Embrapii'])
    adicionar_campo('Valor Contrapartida Empresa (R$)', linha['Fin_Valor Contrapartida Empresa'])
    adicionar_campo('Valor Aporte Unidade (R$)', linha['Fin_Valor Aporte Unidade'])
    adicionar_campo('Valor Total (R$)', linha['Fin_Valor total'])
    
    # 5. Aspectos Técnicos
    doc.add_heading('5. Aspectos Técnicos', level=2)
    adicionar_campo('TRL Inicial', linha['TRL Inicial'])
    adicionar_campo('TRL Final', linha['TRL Final'])
    adicionar_campo('Entregável', linha['Entregável'])
    adicionar_campo('Área de Aplicação', linha['Classificação Área de Aplicação_nível e subnível'])
    adicionar_campo('Tecnologia Habilitadora', linha['Classificação_Tecnologia Habilitadora_nível e subnível'])
    adicionar_campo('Tecnologia Verde', linha['Classificação_Tecnologia Verde_nível e subnível'])
    adicionar_campo('Número de Pedidos de PI', linha['Número de Pedidos de PI'])
    
    # Rodapé Padrão
    doc.add_paragraph('\nDocumento gerado automaticamente a partir da base de projetos.')
    
    return doc

def automatizar_geracao_fichas(caminho_planilha):
    """
    Função principal que lê a base de dados, seleciona os primeiros projetos 
    e orquestra a geração e salvamento dos arquivos Word.
    """
    try:
        print("Lendo a base de projetos...")
        # Lendo o arquivo Excel fornecido
        df_projetos = pd.read_excel(caminho_planilha)
        
        # Selecionando os 5 primeiros projetos
        projetos_selecionados = df_projetos.head(5)
        
        print(f"{len(projetos_selecionados)} projetos selecionados. Iniciando a geração das fichas...")
        
        # Iteração sobre os projetos selecionados
        for index, linha in projetos_selecionados.iterrows():
            codigo_projeto = linha['Projeto_cód']
            nome_arquivo = f"Ficha_Projeto_{codigo_projeto}.docx"
            
            # Chama a função de criação do documento
            doc = criar_ficha_projeto(linha)
            
            # Salva o arquivo gerado
            doc.save(nome_arquivo)
            print(f"Sucesso: {nome_arquivo} gerado!")
            
        print("\nProcesso de automação concluído com sucesso!")
        
    except FileNotFoundError:
        print(f"Erro: O arquivo '{caminho_planilha}' não foi encontrado.")
    except Exception as e:
        print(f"Erro inesperado ao processar os dados: {e}")

if __name__ == "__main__":
    # Nome do arquivo de base passado pela Embrapii
    arquivo_base = 'Embrapii_seleção_analista_2026_questao04_Base_Automação.xlsx'
    automatizar_geracao_fichas(arquivo_base)